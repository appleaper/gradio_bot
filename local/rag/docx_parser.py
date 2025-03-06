
from docx import Document
import re
import pandas as pd
from collections import Counter
from local.rag import rag_tokenizer
from io import BytesIO

'''
pip install python-docx
在 Python 中方便地创建、读取和修改 Word 文档
'''


class RAGFlowDocxParser:

    def __extract_table_content(self, tb):
        # print('deepdoc/parser/docx_parser.py/RAGFlowDocxParser/__extract_table_content')
        '''
        从给定的 tb（可能是一个表格对象）中提取内容。
        '''
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):
        # print('deepdoc/parser/docx_parser.py/RAGFlowDocxParser/__compose_table_content')
        '''
        对提取的表格内容（DataFrame 对象 df）进行处理和组合，生成最终的表格内容。
        '''
        def blockType(b):
            patt = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]
            for p, n in patt:
                if re.search(p, b):
                    return n
            tks = [t for t in rag_tokenizer.tokenize(b).split() if len(t) > 1]
            if len(tks) > 3:
                if len(tks) < 12:
                    return "Tx"
                else:
                    return "Lx"

            if len(tks) == 1 and rag_tokenizer.tag(tks[0]) == "nr":
                return "Nr"

            return "Ot"

        if len(df) < 2:
            return []
        max_type = Counter([blockType(str(df.iloc[i, j])) for i in range(
            1, len(df)) for j in range(len(df.iloc[i, :]))])
        max_type = max(max_type.items(), key=lambda x: x[1])[0]

        colnm = len(df.iloc[0, :])
        hdrows = [0]  # header is not nessesarily appear in the first line
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j]))
                              for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)

        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            hr = [r - i for r in hdrows]
            hr = [r for r in hr if r < 0]
            t = len(hr) - 1
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
            headers = []
            for j in range(len(df.iloc[i, :])):
                t = []
                for h in hr:
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
            cells = []
            for j in range(len(df.iloc[i, :])):
                if not str(df.iloc[i, j]):
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))

        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=100000000):
        # print('deepdoc/parser/docx_parser.py/RAGFlowDocxParser/__call__')
        '''
        将类实例作为可调用对象，用于处理给定的文档（fnm 可以是文档文件名或字节流），并提取指定页面范围
        （从 from_page 到 to_page）的段落和表格内容。
        '''
        self.doc = Document(fnm) if isinstance(
            fnm, str) else Document(BytesIO(fnm))
        pn = 0 # parsed page
        secs = [] # parsed contents
        for p in self.doc.paragraphs:
            if pn > to_page:
                break

            runs_within_single_paragraph = [] # save runs within the range of pages
            for run in p.runs:
                if pn > to_page:
                    break
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text) # append run.text first

                # wrap page break checker into a static method
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1

            # secs.append(("".join(runs_within_single_paragraph), p.style.name if hasattr(p.style, 'name') else '')) # then concat run.text as part of the paragraph
            secs.append("".join(runs_within_single_paragraph))

        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        tbls = ['\n'.join(i) for i in tbls]
        return secs, tbls

if __name__ == "__main__":
    parser = RAGFlowDocxParser()
    filename = "/home/pandas/snap/code/ragflow-new/sdk/python/test/test_sdk_api/test_data/test.docx"  # 替换为你的 Word 文档文件名
    sections, tables = parser(filename)
    # sections返回的是列表，列表中每个元素是一个元祖，每个元祖有2个元素，第一个是文字，第二个是层级比如说heading，Normal
    print("Sections:", sections)
    print("Tables:", tables)